from docx import Document
from docx.shared import Pt
import os

def generate_pdf_report(requirements_paragraph: str, test_cases: str, output_dir="outputs"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()

    # Title
    doc.add_heading('Code analysis report', 0)

    # Section: Business Requirements
    doc.add_heading('Business Requirements', level=1)
    para = doc.add_paragraph(requirements_paragraph)
    para.style.font.size = Pt(12)

    # Section: Test Cases
    doc.add_heading('Generated Test Cases', level=1)
    test_para = doc.add_paragraph(test_cases)
    test_para.style.font.size = Pt(11)

    # Save as Word doc
    docx_path = os.path.join(output_dir, "stakeholder_report.docx")
    doc.save(docx_path)

    # Convert to PDF (if on Windows or macOS)
    try:
        from docx2pdf import convert
        convert(docx_path)
        print(f"✅ PDF report generated at: {docx_path.replace('.docx', '.pdf')}")
    except Exception as e:
        print("⚠️ PDF conversion failed. Word (.docx) file is ready instead.")
        print(str(e))
        return docx_path

    return docx_path.replace(".docx", ".pdf")
